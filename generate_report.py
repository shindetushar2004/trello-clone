#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_with_style(doc, text, level, color=(37, 99, 235)):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
    return heading

def add_colored_paragraph(doc, text, bold=False, color=(31, 41, 55)):
    """Add a paragraph with specific color"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        if bold:
            run.bold = True
    return p

def shade_cell(cell, color):
    """Shade table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create Document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# =================== TITLE PAGE ===================
title = doc.add_heading('IMPLEMENTATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(37, 99, 235)
    run.font.size = Pt(28)
    run.bold = True

doc.add_paragraph()
subtitle = doc.add_heading('Trello Clone Full-Stack Application', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

# Project Info Table
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'

cells = table.rows[0].cells
cells[0].text = 'Project Name'
cells[1].text = 'Trello Clone'
cells[0].paragraphs[0].runs[0].bold = True

cells = table.rows[1].cells
cells[0].text = 'Version'
cells[1].text = '0.1.0'
cells[0].paragraphs[0].runs[0].bold = True

cells = table.rows[2].cells
cells[0].text = 'Date'
cells[1].text = datetime.now().strftime('%B %d, %Y')
cells[0].paragraphs[0].runs[0].bold = True

cells = table.rows[3].cells
cells[0].text = 'Status'
cells[1].text = 'In Development'
cells[0].paragraphs[0].runs[0].bold = True

cells = table.rows[4].cells
cells[0].text = 'Framework'
cells[1].text = 'Next.js 15.3.4 with TypeScript'
cells[0].paragraphs[0].runs[0].bold = True

cells = table.rows[5].cells
cells[0].text = 'Database'
cells[1].text = 'Supabase (PostgreSQL)'
cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# =================== TABLE OF CONTENTS ===================
add_heading_with_style(doc, 'TABLE OF CONTENTS', 1)
toc_items = [
    '1. Project Overview',
    '2. System Architecture',
    '3. Database Schema',
    '4. API Endpoints',
    '5. Features & Functionality',
    '6. Setup Instructions',
    '7. Tech Stack Details',
    '8. File Structure'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# =================== 1. PROJECT OVERVIEW ===================
add_heading_with_style(doc, '1. PROJECT OVERVIEW', 1)

doc.add_heading('1.1 Executive Summary', level=2)
doc.add_paragraph(
    'Trello Clone is a full-stack web application built with Next.js and Supabase that replicates '
    'the core functionality of Trello. The application enables users to create boards, organize tasks '
    'into columns, manage team collaboration through invitations, and leverage real-time features for '
    'seamless workflow management.'
)

doc.add_heading('1.2 Project Goals', level=2)
goals = [
    'Create a production-ready task management application',
    'Implement real-time data synchronization',
    'Support multi-user collaboration with invite system',
    'Provide drag-and-drop task management interface',
    'Ensure secure authentication and authorization',
    'Build scalable architecture for future enhancements'
]
for goal in goals:
    doc.add_paragraph(goal, style='List Bullet')

doc.add_heading('1.3 Key Features', level=2)
features = [
    'Board Management - Create, update, and manage multiple boards',
    'Column Organization - Organize tasks into custom columns',
    'Task Management - Create, edit, delete, and prioritize tasks',
    'Team Collaboration - Invite team members to boards via email',
    'Authentication - Secure user authentication via Clerk',
    'Real-time Updates - Real-time data synchronization with Supabase',
    'Responsive Design - Mobile-friendly UI built with Tailwind CSS',
    'Drag & Drop - Smooth drag-and-drop interface using dnd-kit'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# =================== 2. SYSTEM ARCHITECTURE ===================
add_heading_with_style(doc, '2. SYSTEM ARCHITECTURE', 1)

doc.add_heading('2.1 Architecture Overview', level=2)
doc.add_paragraph(
    'The application follows a modern client-server architecture with a Next.js frontend, '
    'API layer, and Supabase backend. The system is designed to be scalable, maintainable, '
    'and follows best practices for full-stack development.'
)

doc.add_heading('2.2 Architectural Layers', level=2)

# Create architecture table
arch_table = doc.add_table(rows=4, cols=3)
arch_table.style = 'Light Grid Accent 1'

header_cells = arch_table.rows[0].cells
header_cells[0].text = 'Layer'
header_cells[1].text = 'Technology'
header_cells[2].text = 'Responsibility'
for cell in header_cells:
    shade_cell(cell, '2563EB')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

arch_table.rows[1].cells[0].text = 'Frontend'
arch_table.rows[1].cells[1].text = 'Next.js, React, Tailwind CSS'
arch_table.rows[1].cells[2].text = 'User interface, state management, real-time updates'

arch_table.rows[2].cells[0].text = 'API Layer'
arch_table.rows[2].cells[1].text = 'Next.js API Routes'
arch_table.rows[2].cells[2].text = 'Business logic, authentication, email notifications'

arch_table.rows[3].cells[0].text = 'Backend'
arch_table.rows[3].cells[1].text = 'Supabase (PostgreSQL)'
arch_table.rows[3].cells[2].text = 'Data persistence, real-time subscriptions'

doc.add_page_break()

# =================== 3. DATABASE SCHEMA ===================
add_heading_with_style(doc, '3. DATABASE SCHEMA', 1)

doc.add_heading('3.1 Core Tables', level=2)

doc.add_heading('3.1.1 Boards Table', level=3)
boards_table = doc.add_table(rows=8, cols=3)
boards_table.style = 'Light Grid Accent 1'

header = boards_table.rows[0].cells
header[0].text = 'Column'
header[1].text = 'Type'
header[2].text = 'Description'
for cell in header:
    shade_cell(cell, '2563EB')
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

boards_data = [
    ('id', 'UUID', 'Primary key, unique board identifier'),
    ('title', 'VARCHAR', 'Board title/name'),
    ('description', 'TEXT', 'Optional board description'),
    ('color', 'VARCHAR', 'Board color/theme'),
    ('user_id', 'UUID', 'Owner of the board'),
    ('created_at', 'TIMESTAMP', 'Board creation timestamp'),
    ('updated_at', 'TIMESTAMP', 'Last update timestamp')
]

for i, (col, type_, desc) in enumerate(boards_data, 1):
    boards_table.rows[i].cells[0].text = col
    boards_table.rows[i].cells[1].text = type_
    boards_table.rows[i].cells[2].text = desc

doc.add_heading('3.1.2 Columns Table', level=3)
columns_table = doc.add_table(rows=7, cols=3)
columns_table.style = 'Light Grid Accent 1'

header = columns_table.rows[0].cells
header[0].text = 'Column'
header[1].text = 'Type'
header[2].text = 'Description'
for cell in header:
    shade_cell(cell, '2563EB')
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

columns_data = [
    ('id', 'UUID', 'Primary key'),
    ('board_id', 'UUID', 'Foreign key to boards table'),
    ('title', 'VARCHAR', 'Column title (e.g., "To Do", "In Progress")'),
    ('sort_order', 'INTEGER', 'Column order on the board'),
    ('created_at', 'TIMESTAMP', 'Creation timestamp'),
    ('user_id', 'UUID', 'Owner reference')
]

for i, (col, type_, desc) in enumerate(columns_data, 1):
    columns_table.rows[i].cells[0].text = col
    columns_table.rows[i].cells[1].text = type_
    columns_table.rows[i].cells[2].text = desc

doc.add_heading('3.1.3 Tasks Table', level=3)
tasks_table = doc.add_table(rows=10, cols=3)
tasks_table.style = 'Light Grid Accent 1'

header = tasks_table.rows[0].cells
header[0].text = 'Column'
header[1].text = 'Type'
header[2].text = 'Description'
for cell in header:
    shade_cell(cell, '2563EB')
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

tasks_data = [
    ('id', 'UUID', 'Primary key'),
    ('column_id', 'UUID', 'Foreign key to columns table'),
    ('title', 'VARCHAR', 'Task title'),
    ('description', 'TEXT', 'Detailed task description'),
    ('assignee', 'VARCHAR', 'Assigned team member'),
    ('due_date', 'DATE', 'Task deadline'),
    ('priority', 'ENUM', 'Priority level: low, medium, high'),
    ('sort_order', 'INTEGER', 'Task order within column'),
    ('created_at', 'TIMESTAMP', 'Creation timestamp')
]

for i, (col, type_, desc) in enumerate(tasks_data, 1):
    tasks_table.rows[i].cells[0].text = col
    tasks_table.rows[i].cells[1].text = type_
    tasks_table.rows[i].cells[2].text = desc

doc.add_heading('3.1.4 Board Invites Table', level=3)
invites_table = doc.add_table(rows=7, cols=3)
invites_table.style = 'Light Grid Accent 1'

header = invites_table.rows[0].cells
header[0].text = 'Column'
header[1].text = 'Type'
header[2].text = 'Description'
for cell in header:
    shade_cell(cell, '2563EB')
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

invites_data = [
    ('id', 'UUID', 'Primary key'),
    ('board_id', 'UUID', 'Foreign key to boards table'),
    ('invited_email', 'VARCHAR', 'Email of invited user'),
    ('token', 'VARCHAR', 'Unique invite token'),
    ('invited_by', 'UUID', 'User who sent the invite'),
    ('created_at', 'TIMESTAMP', 'Invite creation time')
]

for i, (col, type_, desc) in enumerate(invites_data, 1):
    invites_table.rows[i].cells[0].text = col
    invites_table.rows[i].cells[1].text = type_
    invites_table.rows[i].cells[2].text = desc

doc.add_page_break()

# =================== 4. API ENDPOINTS ===================
add_heading_with_style(doc, '4. API ENDPOINTS', 1)

doc.add_heading('4.1 Authentication & Invites', level=2)

# Invite API endpoint
doc.add_heading('4.1.1 POST /api/invite', level=3)
doc.add_paragraph('Purpose: Send board invitation to team member', style='List Bullet')
doc.add_paragraph('Request Body:', style='List Bullet')
request_body = doc.add_paragraph('', style='List Bullet 2')
request_body.text = '{  boardId: string,  email: string,  boardTitle: string }'

doc.add_paragraph('Response:', style='List Bullet')
response_p = doc.add_paragraph('', style='List Bullet 2')
response_p.text = '{ success: boolean, message: string }'

doc.add_paragraph('Status Codes:', style='List Bullet')
doc.add_paragraph('200: Success', style='List Bullet 2')
doc.add_paragraph('401: Unauthorized', style='List Bullet 2')
doc.add_paragraph('400: Missing fields', style='List Bullet 2')
doc.add_paragraph('500: Server error', style='List Bullet 2')

# Accept Invite
doc.add_heading('4.1.2 POST /api/invite/accept', level=3)
doc.add_paragraph('Purpose: Accept a board invitation', style='List Bullet')
doc.add_paragraph('Request Body:', style='List Bullet')
accept_body = doc.add_paragraph('', style='List Bullet 2')
accept_body.text = '{ token: string }'

doc.add_paragraph('Response:', style='List Bullet')
response_p = doc.add_paragraph('', style='List Bullet 2')
response_p.text = '{ success: boolean, boardId: string }'

# Get Invite
doc.add_heading('4.1.3 GET /api/invite/token/[token]', level=3)
doc.add_paragraph('Purpose: Retrieve invite details before acceptance', style='List Bullet')
doc.add_paragraph('Parameters:', style='List Bullet')
doc.add_paragraph('token: string (URL parameter)', style='List Bullet 2')
doc.add_paragraph('Response:', style='List Bullet')
resp_p = doc.add_paragraph('', style='List Bullet 2')
resp_p.text = '{ boardId: string, boardTitle: string, invitedEmail: string }'

doc.add_page_break()

# =================== 5. FEATURES & FUNCTIONALITY ===================
add_heading_with_style(doc, '5. FEATURES & FUNCTIONALITY', 1)

doc.add_heading('5.1 Board Management', level=2)
doc.add_paragraph('Users can create, view, update, and delete boards with the following capabilities:')
board_features = [
    'Create new boards with custom titles and descriptions',
    'Set board color/theme for visual organization',
    'View all personal boards in dashboard',
    'Update board details',
    'Delete boards (owner only)',
    'Share boards via email invitations'
]
for feature in board_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('5.2 Task Organization', level=2)
doc.add_paragraph('Organize tasks within columns for workflow management:')
task_features = [
    'Create columns with custom titles (e.g., "To Do", "In Progress", "Done")',
    'Reorder columns via drag-and-drop',
    'Create tasks within columns',
    'Set task priority levels (Low, Medium, High)',
    'Add task descriptions and due dates',
    'Assign tasks to team members',
    'Drag tasks between columns',
    'Delete tasks'
]
for feature in task_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('5.3 Team Collaboration', level=2)
doc.add_paragraph('Collaborate with team members through invitation system:')
collab_features = [
    'Send board invitations via email',
    'Unique invite tokens with 7-day expiration',
    'Accept/decline invitations',
    'View board members',
    'Share access to specific boards',
    'Email notifications for invitations'
]
for feature in collab_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('5.4 Real-time Features', level=2)
doc.add_paragraph('Supabase real-time subscriptions enable:')
realtime_features = [
    'Live task updates across devices',
    'Real-time column modifications',
    'Instant board changes for all members',
    'Automatic UI sync without page refresh'
]
for feature in realtime_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('5.5 User Authentication', level=2)
doc.add_paragraph('Secure authentication via Clerk:')
auth_features = [
    'Sign up with email',
    'Social login (Google, GitHub, etc.)',
    'Session management',
    'Protected routes and API endpoints',
    'User profile management'
]
for feature in auth_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# =================== 6. SETUP INSTRUCTIONS ===================
add_heading_with_style(doc, '6. SETUP INSTRUCTIONS', 1)

doc.add_heading('6.1 Prerequisites', level=2)
doc.add_paragraph('Before starting, ensure you have:')
prerequisites = [
    'Node.js 18+ and npm/yarn',
    'Supabase account and project',
    'Clerk account for authentication',
    'Resend account for email notifications',
    'Git for version control'
]
for prereq in prerequisites:
    doc.add_paragraph(prereq, style='List Bullet')

doc.add_heading('6.2 Environment Setup', level=2)
doc.add_paragraph('1. Clone the repository:')
code_p = doc.add_paragraph('git clone <repository-url>', style='List Bullet')
code_p.paragraph_format.left_indent = Inches(0.5)

doc.add_paragraph('2. Install dependencies:')
code_p = doc.add_paragraph('npm install', style='List Bullet')
code_p.paragraph_format.left_indent = Inches(0.5)

doc.add_paragraph('3. Create .env.local file with:')
env_vars = [
    'NEXT_PUBLIC_SUPABASE_URL=your_supabase_url',
    'NEXT_PUBLIC_SUPABASE_ANON_KEY=your_anon_key',
    'SUPABASE_SERVICE_ROLE_KEY=your_service_role_key',
    'NEXT_PUBLIC_CLERK_PUBLISHABLE_KEY=your_clerk_key',
    'CLERK_SECRET_KEY=your_clerk_secret',
    'RESEND_API_KEY=your_resend_key',
    'NEXT_PUBLIC_APP_URL=http://localhost:3000'
]
for var in env_vars:
    doc.add_paragraph(var, style='List Bullet')

doc.add_heading('6.3 Database Setup', level=2)
doc.add_paragraph('1. Create tables in Supabase SQL Editor:')
doc.add_paragraph('Execute the schema migration scripts to create:')
db_tables = [
    'boards table',
    'columns table',
    'tasks table',
    'board_invites table'
]
for table in db_tables:
    doc.add_paragraph(table, style='List Bullet')

doc.add_paragraph('2. Enable Row-Level Security (RLS) on all tables')
doc.add_paragraph('3. Configure RLS policies for user isolation')

doc.add_heading('6.4 Running the Application', level=2)
doc.add_paragraph('1. Start development server:')
code_p = doc.add_paragraph('npm run dev', style='List Bullet')
code_p.paragraph_format.left_indent = Inches(0.5)

doc.add_paragraph('2. Open browser and navigate to http://localhost:3000')
doc.add_paragraph('3. Sign up with Clerk authentication')
doc.add_paragraph('4. Create your first board')

doc.add_heading('6.5 Deployment', level=2)
doc.add_paragraph('Deploy to Vercel:')
deploy_steps = [
    'Push code to GitHub repository',
    'Connect repository to Vercel',
    'Add environment variables in Vercel dashboard',
    'Deploy with a single click'
]
for step in deploy_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_page_break()

# =================== 7. TECH STACK DETAILS ===================
add_heading_with_style(doc, '7. TECH STACK DETAILS', 1)

doc.add_heading('7.1 Frontend Technologies', level=2)

tech_table = doc.add_table(rows=9, cols=3)
tech_table.style = 'Light Grid Accent 1'

header = tech_table.rows[0].cells
header[0].text = 'Technology'
header[1].text = 'Version'
header[2].text = 'Purpose'
for cell in header:
    shade_cell(cell, '2563EB')
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

tech_data = [
    ('Next.js', '15.3.4', 'React framework for production-ready applications'),
    ('React', '19.0.0', 'UI library for building components'),
    ('TypeScript', '5.0', 'Type safety and better development experience'),
    ('Tailwind CSS', '4.0', 'Utility-first CSS framework'),
    ('@dnd-kit', '6.3.1+', 'Drag and drop functionality'),
    ('@clerk/nextjs', '6.24.0', 'Authentication and user management'),
    ('@supabase/ssr', '0.6.1', 'Supabase integration for server-side rendering'),
    ('Lucide React', '0.525.0', 'Icon library for UI components')
]

for i, (tech, version, purpose) in enumerate(tech_data, 1):
    tech_table.rows[i].cells[0].text = tech
    tech_table.rows[i].cells[1].text = version
    tech_table.rows[i].cells[2].text = purpose

doc.add_heading('7.2 Backend Services', level=2)

doc.add_paragraph('Supabase: Hosted PostgreSQL database')
doc.add_paragraph('Features:', style='List Bullet')
supabase_features = [
    'Real-time subscriptions',
    'Row-level security (RLS)',
    'PostgreSQL database',
    'RESTful API',
    'Authentication integration'
]
for feature in supabase_features:
    doc.add_paragraph(feature, style='List Bullet 2')

doc.add_paragraph('Clerk: Authentication & User Management')
doc.add_paragraph('Features:', style='List Bullet')
clerk_features = [
    'Email/password authentication',
    'Social login providers',
    'Session management',
    'User profiles',
    'Multi-tenancy support'
]
for feature in clerk_features:
    doc.add_paragraph(feature, style='List Bullet 2')

doc.add_paragraph('Resend: Email Notifications')
doc.add_paragraph('Features:', style='List Bullet')
resend_features = [
    'Send invitation emails',
    'Templated emails',
    'Delivery tracking',
    'Reliable email service'
]
for feature in resend_features:
    doc.add_paragraph(feature, style='List Bullet 2')

doc.add_page_break()

# =================== 8. FILE STRUCTURE ===================
add_heading_with_style(doc, '8. FILE STRUCTURE', 1)

doc.add_heading('8.1 Project Directory Layout', level=2)

structure_text = '''
trello-clone-fullstack/
├── app/                           # Next.js app directory
│   ├── api/                       # API routes
│   │   └── invite/
│   │       ├── route.ts          # Send invitation endpoint
│   │       └── accept/
│   │           └── route.ts      # Accept invitation endpoint
│   ├── boards/
│   │   └── [id]/                 # Board detail page
│   │       └── page.tsx
│   ├── dashboard/                 # User dashboard
│   │   ├── layout.tsx
│   │   └── page.tsx
│   ├── pricing/                   # Pricing page
│   │   └── page.tsx
│   ├── globals.css               # Global styles
│   ├── layout.tsx                # Root layout
│   └── page.tsx                  # Home page
├── components/                    # React components
│   ├── navbar.tsx                # Navigation bar
│   └── ui/                       # UI component library
│       ├── badge.tsx
│       ├── button.tsx
│       ├── card.tsx
│       ├── dialog.tsx
│       ├── input.tsx
│       ├── label.tsx
│       ├── select.tsx
│       └── textarea.tsx
├── lib/                          # Utilities and helpers
│   ├── services.ts              # Business logic
│   ├── utils.ts                 # Utility functions
│   ├── contexts/
│   │   └── PlanContext.tsx      # Plan/pricing context
│   ├── hooks/
│   │   └── useBoards.ts         # Board data hook
│   └── supabase/
│       ├── models.ts            # TypeScript interfaces
│       ├── server.ts            # Server-side Supabase client
│       └── SupabaseProvider.tsx # Client provider
├── public/                       # Static assets
├── package.json                 # Dependencies
├── tsconfig.json               # TypeScript config
├── next.config.ts              # Next.js config
├── middleware.ts               # Next.js middleware
├── tailwind.config.ts          # Tailwind CSS config
├── eslint.config.mjs           # ESLint config
└── postcss.config.mjs          # PostCSS config
'''

doc.add_paragraph(structure_text, style='List Paragraph')

doc.add_heading('8.2 Key Files Description', level=2)

files_table = doc.add_table(rows=7, cols=2)
files_table.style = 'Light Grid Accent 1'

header = files_table.rows[0].cells
header[0].text = 'File'
header[1].text = 'Description'
for cell in header:
    shade_cell(cell, '2563EB')
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

files_data = [
    ('app/layout.tsx', 'Root layout with provider setup'),
    ('lib/supabase/models.ts', 'TypeScript type definitions for database'),
    ('lib/supabase/server.ts', 'Server-side Supabase client configuration'),
    ('lib/hooks/useBoards.ts', 'Custom React hook for board data'),
    ('app/api/invite/route.ts', 'Email invitation API endpoint'),
    ('middleware.ts', 'Next.js middleware for routing/auth')
]

for i, (file, desc) in enumerate(files_data, 1):
    files_table.rows[i].cells[0].text = file
    files_table.rows[i].cells[1].text = desc

doc.add_page_break()

# =================== CONCLUSION ===================
add_heading_with_style(doc, 'CONCLUSION', 1)

doc.add_paragraph(
    'The Trello Clone Full-Stack Application represents a modern, scalable solution for task '
    'management and team collaboration. Built with cutting-edge technologies including Next.js, '
    'Supabase, and Clerk, the application provides a robust foundation for real-time collaboration features.'
)

doc.add_heading('Key Achievements:', level=2)
achievements = [
    'Production-ready code structure',
    'Real-time data synchronization',
    'Secure authentication system',
    'Email-based team collaboration',
    'Responsive and user-friendly interface',
    'Scalable database architecture'
]
for achievement in achievements:
    doc.add_paragraph(achievement, style='List Bullet')

doc.add_heading('Future Enhancements:', level=2)
future = [
    'Premium subscription tiers',
    'Advanced task analytics',
    'Mobile app development',
    'Integration with third-party tools',
    'Advanced permission controls',
    'Activity logging and audit trails'
]
for item in future:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# =================== APPENDIX ===================
doc.add_heading('APPENDIX', 1)

doc.add_heading('A. Common Environment Variables', level=2)
env_table = doc.add_table(rows=8, cols=2)
env_table.style = 'Light Grid Accent 1'

env_header = env_table.rows[0].cells
env_header[0].text = 'Variable'
env_header[1].text = 'Description'
for cell in env_header:
    shade_cell(cell, '2563EB')
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

env_vars_data = [
    ('NEXT_PUBLIC_SUPABASE_URL', 'Supabase project URL'),
    ('NEXT_PUBLIC_SUPABASE_ANON_KEY', 'Supabase anonymous key for client-side'),
    ('SUPABASE_SERVICE_ROLE_KEY', 'Service role key for server-side admin operations'),
    ('NEXT_PUBLIC_CLERK_PUBLISHABLE_KEY', 'Clerk public key for authentication'),
    ('CLERK_SECRET_KEY', 'Clerk secret key for server-side operations'),
    ('RESEND_API_KEY', 'Resend API key for sending emails'),
    ('NEXT_PUBLIC_APP_URL', 'Application URL for invite links')
]

for i, (var, desc) in enumerate(env_vars_data, 1):
    env_table.rows[i].cells[0].text = var
    env_table.rows[i].cells[1].text = desc

doc.add_heading('B. Development Commands', level=2)
commands = [
    'npm run dev - Start development server with Turbopack',
    'npm run build - Build production bundle',
    'npm start - Start production server',
    'npm run lint - Run ESLint for code quality'
]
for cmd in commands:
    doc.add_paragraph(cmd, style='List Bullet')

doc.add_heading('C. Troubleshooting', level=2)
doc.add_paragraph('Issue: Database connection errors')
doc.add_paragraph('Solution: Verify Supabase credentials in .env.local', style='List Bullet')

doc.add_paragraph('Issue: Authentication not working')
doc.add_paragraph('Solution: Check Clerk keys and ClerkProvider setup', style='List Bullet')

doc.add_paragraph('Issue: Emails not sending')
doc.add_paragraph('Solution: Verify Resend API key and email permissions', style='List Bullet')

# Save document
doc.save('report.docx')
print("Report generated successfully: report.docx")
